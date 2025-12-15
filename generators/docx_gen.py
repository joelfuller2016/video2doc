"""
DOCX Generator - Creates Microsoft Word documents from video analysis
Produces professional instruction manuals with embedded screenshots
"""

import os
import sys
from pathlib import Path
from typing import List, Dict, Optional
from datetime import datetime

# Add parent directory to path for imports when running as module
sys.path.insert(0, str(Path(__file__).parent.parent))
from logger import get_logger
from utils.path_validator import (
    is_safe_path,
    validate_file_output_path,
)

# Module logger
logger = get_logger(__name__)

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE


def create_styles(doc: Document) -> None:
    """Add custom styles to the document"""
    styles = doc.styles

    # Title style
    if "Doc Title" not in [s.name for s in styles]:
        title_style = styles.add_style("Doc Title", WD_STYLE_TYPE.PARAGRAPH)
        title_style.font.size = Pt(28)
        title_style.font.bold = True
        title_style.font.color.rgb = RGBColor(0, 51, 102)

    # Section Header style
    if "Section Header" not in [s.name for s in styles]:
        section_style = styles.add_style("Section Header", WD_STYLE_TYPE.PARAGRAPH)
        section_style.font.size = Pt(16)
        section_style.font.bold = True
        section_style.font.color.rgb = RGBColor(0, 102, 153)


def generate_docx(
    title: str,
    summary: str,
    sections: List[Dict],
    screenshots: Dict[float, str],
    output_path: str,
    author: str = "FrameNotes",
    image_width: float = 6.0
) -> str:
    """
    Generate a Microsoft Word document from video analysis.

    Args:
        title: Document title
        summary: Brief summary/introduction
        sections: List of section dictionaries with title, content, screenshot_timestamp
        screenshots: Dict mapping timestamps to image file paths
        output_path: Path for output .docx file
        author: Document author name
        image_width: Width of images in inches

    Returns:
        Path to generated document
    """
    logger.info(f"Generating DOCX document: {output_path}")
    logger.debug(f"Document has {len(sections)} sections, {len(screenshots)} screenshots")

    # Validate output path for security
    output_result = validate_file_output_path(
        output_path,
        allowed_extensions=[".docx"],
        create_parent_dirs=True
    )
    if not output_result.is_valid:
        raise ValueError(f"Invalid output path: {output_result.error_message}")
    output_path = output_result.sanitized_value

    # Validate screenshot paths for security
    validated_screenshots = {}
    for ts, screenshot_path in screenshots.items():
        is_safe, error = is_safe_path(screenshot_path)
        if is_safe and os.path.exists(screenshot_path):
            validated_screenshots[ts] = screenshot_path
        else:
            logger.warning(f"Skipping invalid screenshot path at {ts}: {error or 'File not found'}")
    screenshots = validated_screenshots

    doc = Document()

    # Set document properties
    doc.core_properties.author = author
    doc.core_properties.title = title
    doc.core_properties.created = datetime.now()

    # Create custom styles
    create_styles(doc)

    # Add title
    title_para = doc.add_paragraph()
    title_run = title_para.add_run(title)
    title_run.font.size = Pt(28)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0, 51, 102)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add generation date
    date_para = doc.add_paragraph()
    date_run = date_para.add_run(f"Generated: {datetime.now().strftime('%B %d, %Y')}")
    date_run.font.size = Pt(10)
    date_run.font.italic = True
    date_run.font.color.rgb = RGBColor(128, 128, 128)
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()  # Spacer

    # Add summary/introduction
    if summary:
        intro_heading = doc.add_heading("Overview", level=1)
        summary_para = doc.add_paragraph(summary)
        doc.add_paragraph()  # Spacer

    # Add table of contents placeholder
    toc_heading = doc.add_heading("Table of Contents", level=1)
    for i, section in enumerate(sections, 1):
        toc_item = doc.add_paragraph()
        toc_item.add_run(f"{i}. {section.get('title', 'Section')}")

    doc.add_page_break()

    # Add sections
    logger.debug("Adding document sections...")
    for i, section in enumerate(sections, 1):
        section_title = section.get("title", f"Section {i}")
        section_content = section.get("content", "")
        screenshot_ts = section.get("screenshot_timestamp")
        screenshot_reason = section.get("screenshot_reason", "")

        # Section heading
        heading = doc.add_heading(f"{i}. {section_title}", level=1)

        # Section content
        if section_content:
            content_para = doc.add_paragraph(section_content)

        # Add screenshot if available
        if screenshot_ts is not None:
            # Find matching screenshot
            screenshot_path = None

            # Look for exact match first
            if screenshot_ts in screenshots:
                screenshot_path = screenshots[screenshot_ts]
            else:
                # Find closest timestamp
                closest_ts = min(screenshots.keys(), key=lambda x: abs(x - screenshot_ts), default=None)
                if closest_ts is not None and abs(closest_ts - screenshot_ts) < 5:  # Within 5 seconds
                    screenshot_path = screenshots[closest_ts]

            if screenshot_path and os.path.exists(screenshot_path):
                try:
                    logger.debug(f"Embedding screenshot for section {i}: {screenshot_path}")
                    # Add image
                    doc.add_paragraph()  # Spacer before image

                    # Center the image
                    img_para = doc.add_paragraph()
                    img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = img_para.add_run()
                    run.add_picture(screenshot_path, width=Inches(image_width))

                    # Add caption
                    if screenshot_reason:
                        caption = doc.add_paragraph()
                        caption_run = caption.add_run(f"Figure {i}: {screenshot_reason}")
                        caption_run.font.size = Pt(10)
                        caption_run.font.italic = True
                        caption_run.font.color.rgb = RGBColor(100, 100, 100)
                        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER

                    doc.add_paragraph()  # Spacer after image

                except Exception as e:
                    logger.warning(f"Failed to embed image for section {i}: {e}. Skipping image.")

        # Add some spacing between sections
        doc.add_paragraph()

    # Add footer with generation info
    doc.add_page_break()
    footer_heading = doc.add_heading("Document Information", level=1)

    info_para = doc.add_paragraph()
    info_para.add_run("Generated by: ").bold = True
    info_para.add_run("FrameNotes - AI-Powered Video Documentation\n")
    info_para.add_run("Generation Date: ").bold = True
    info_para.add_run(f"{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
    info_para.add_run("Total Sections: ").bold = True
    info_para.add_run(f"{len(sections)}\n")
    info_para.add_run("Total Screenshots: ").bold = True
    info_para.add_run(f"{len(screenshots)}")

    # Ensure output directory exists
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    # Save document
    doc.save(str(output_path))
    logger.info(f"DOCX document saved: {output_path}")

    return str(output_path)


if __name__ == "__main__":
    from logger import init_logging
    init_logging(verbose=True)

    # Test with mock data
    test_sections = [
        {
            "title": "Introduction",
            "content": "This document provides a comprehensive guide to the topic covered in the video.",
            "screenshot_timestamp": 5.0,
            "screenshot_reason": "Opening screen of the application"
        },
        {
            "title": "Getting Started",
            "content": "Follow these steps to begin using the software effectively.",
            "screenshot_timestamp": 30.0,
            "screenshot_reason": "Main interface overview"
        },
        {
            "title": "Advanced Features",
            "content": "Explore the advanced capabilities for power users.",
            "screenshot_timestamp": 60.0,
            "screenshot_reason": "Advanced settings panel"
        }
    ]

    # Generate without actual screenshots for testing
    output = generate_docx(
        title="Test Documentation",
        summary="This is a test document generated from video content.",
        sections=test_sections,
        screenshots={},  # No actual screenshots for test
        output_path="./test_output.docx"
    )

    logger.info(f"Generated: {output}")
